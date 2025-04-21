from docx import Document


def write_to_word(file_path, abook):
    document = Document()
    section = document.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = "十分钟读一本书@🦓"
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = "~ zebura AI studio ~"
    style = document.styles['Normal']
    font = style.font
    font.name = '宋体'

    title = abook['title']
    author = abook['author']
    biography = abook['biography']
    background = abook['background']
    summary = abook['summary']

    document.add_heading(title, level=1)
    document.add_heading(author, level=2)
    document.add_heading("作者小传", level=3)
    document.add_paragraph(biography)
    document.add_heading("创作背景", level=3)
    document.add_paragraph(background)
    document.add_heading("摘要", level=3)
    document.add_paragraph(summary)
    
    quotes = abook['quotes']
    document.add_heading("名言金句", level=3)
    for quote, context in quotes:
        document.add_paragraph(f"- {quote}")
        document.add_paragraph(context)

    document.save(file_path)